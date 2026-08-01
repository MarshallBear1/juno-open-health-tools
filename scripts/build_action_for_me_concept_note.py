#!/usr/bin/env python3
"""Build the Juno x Action for M.E. partnership concept note."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "partnership"
DOCX_PATH = OUT / "juno-action-for-me-data-brief-concept-note.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(27, 38, 52)
MUTED = RGBColor(92, 101, 112)
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=11, color=INK, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def shade_paragraph(paragraph, fill="F4F6F9"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")
    p_pr.append(ind)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([r_pr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_lead(doc, label, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    lead = p.add_run(f"{label}. ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body)
    return p


def add_body(doc, text, bold_prefix=None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run("JUNO x ACTION FOR M.E.  |  CONCEPT NOTE")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("For discussion  |  1 August 2026  |  Page ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(footer)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(6)
    org_run = org.add_run("Juno Open Health Tools")
    set_run_font(org_run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("What gets lost before the appointment?")
    set_run_font(tr, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("A co-authored data brief on communication barriers and low-burden preparation for people with ME/CFS")
    set_run_font(sr, size=13, color=MUTED, italic=True)

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [4680, 4680])
    meta_data = [
        ("Proposed by", "Juno / Juno Open Health Tools"),
        ("Proposed collaborator", "Action for M.E. and the PRIME research involvement community"),
        ("Status", "Scoping proposal - no data collection has begun"),
        ("Suggested first step", "A 30-minute fit and governance conversation"),
    ]
    for idx, (label, value) in enumerate(meta_data):
        row = idx // 2
        col = idx % 2
        cell = meta.cell(row, col)
        shade_cell(cell, "F4F6F9")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=9.5, color=DARK_BLUE, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=9.5, color=INK)
    mark_header_row(meta.rows[0])

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(10)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.25
    shade_paragraph(lead)
    lr = lead.add_run("The proposal. ")
    set_run_font(lr, bold=True, color=DARK_BLUE)
    lbody = lead.add_run(
        "Co-produce a short, public data brief on what people with ME/CFS find hardest to communicate in time-limited appointments, which preparation methods are genuinely usable, and what patients want clinicians to understand before the conversation begins."
    )
    set_run_font(lbody)

    add_heading(doc, "Why this is worth testing", 1)
    add_body(
        doc,
        "Fluctuating symptoms, post-exertional worsening, cognitive load and a changing baseline can be difficult to compress into a few minutes. The result is not simply an information problem: people may leave without having explained the change that mattered most or asked the question they came to ask."
    )
    add_body(
        doc,
        "Action for M.E. places lived experience at the centre of its services and research partnerships. Juno has built an openly licensed, non-diagnostic appointment-preparation method from lived experience of chronic illness. A tightly scoped collaboration could turn those complementary strengths into useful evidence and practical resources without positioning either organisation as a clinical authority."
    )

    add_heading(doc, "Proposed research question", 1)
    add_body(
        doc,
        "What information do people with ME/CFS most struggle to communicate before and during appointments, which low-burden preparation supports help them express changes and priorities more clearly, and what do they want clinicians to ask or acknowledge?",
        italic=True,
    )
    add_lead(doc, "Communication", "Which experiences are hardest to describe, prioritise or put on a short timeline?")
    add_lead(doc, "Usability", "Which preparation formats remain manageable with fatigue, pain, sensory sensitivity or cognitive dysfunction?")
    add_lead(doc, "Equity", "What adjustments are needed for people with severe ME, carers and people whose access needs are often excluded from standard digital tools?")
    doc.add_page_break()

    add_heading(doc, "A small, governance-first study", 1)
    add_body(
        doc,
        "The design should be co-produced and intentionally low burden. Before any recruitment, the partners would decide whether the work is patient and public involvement, service evaluation or research, identify the appropriate ethics and data-protection route, and agree who is the data controller."
    )

    roadmap = doc.add_table(rows=1, cols=3)
    roadmap.style = "Table Grid"
    set_table_geometry(roadmap, [1900, 1700, 5760])
    headers = ["Phase", "Indicative timing", "Decision or output"]
    for idx, text in enumerate(headers):
        cell = roadmap.rows[0].cells[idx]
        shade_cell(cell, "F4F6F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    mark_header_row(roadmap.rows[0])

    phases = [
        ("1. Scope", "Week 1", "Agree the question, classification, governance, access needs and stop/go criteria."),
        ("2. Co-design", "Week 2", "A small lived-experience group reviews language, burden, accessibility and response options."),
        ("3. Listen", "Weeks 3-4", "If approved, run a short anonymous survey and optional low-burden conversations through agreed channels."),
        ("4. Interpret", "Week 5", "Review aggregate findings with lived-experience contributors; explicitly separate observation from inference."),
        ("5. Publish", "Week 6", "Release the co-authored brief, practical conversation guide, methods note and limitations."),
    ]
    for phase, timing, decision in phases:
        cells = roadmap.add_row().cells
        for idx, value in enumerate((phase, timing, decision)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.3, color=INK, bold=(idx == 0))
    set_table_geometry(roadmap, [1900, 1700, 5760])

    add_heading(doc, "Safeguards and ownership", 1)
    add_lead(doc, "No patient records", "The project would not request clinical documents, account credentials, record numbers or information copied from a clinical system.")
    add_lead(doc, "Minimum necessary data", "Collect only what is needed for the agreed question; avoid direct identifiers and free-text prompts that invite unnecessary disclosure.")
    add_lead(doc, "No commercial lead generation", "Participation would not be used to market Juno, build sales lists or target people with advertising.")
    add_lead(doc, "Shared interpretation and veto", "Action for M.E. and lived-experience contributors would help interpret findings and approve co-branding; both partners could stop publication if safety or accuracy concerns remain.")

    add_heading(doc, "Proposed public outputs", 1)
    add_lead(doc, "Data brief", "A concise 4-6 page narrative reporting aggregate themes, response distributions, methods and limitations.")
    add_lead(doc, "Conversation guide", "A one-page resource for patients and clinicians focused on change from baseline, impact, timeline and priorities.")
    add_lead(doc, "Methods appendix", "Question wording, accessibility adaptations, classification decisions and a transparent account of what the study cannot conclude.")
    add_lead(doc, "Open practical template", "An adapted appointment-preparation sheet released under an agreed licence, with no participant-level data published.")

    add_heading(doc, "What each partner could contribute", 1)
    add_lead(doc, "Action for M.E. / PRIME", "Lived-experience governance, research/PPI classification, access and safeguarding advice, contributor relationships, interpretation and publication approval.")
    add_lead(doc, "Juno", "Initial drafting, open templates, accessible prototyping, analysis support under the agreed governance structure, design and publication production.")

    ask = doc.add_paragraph()
    ask.paragraph_format.space_before = Pt(8)
    ask.paragraph_format.space_after = Pt(10)
    ask.paragraph_format.line_spacing = 1.25
    shade_paragraph(ask, "E8EEF5")
    ar = ask.add_run("The immediate ask. ")
    set_run_font(ar, bold=True, color=DARK_BLUE)
    ab = ask.add_run(
        "Would Action for M.E. be open to a 30-minute scoping call to decide whether this question is useful, who should shape it, and what governance route would be proportionate? We are comfortable starting with a very small pilot or deciding together that a different question would be more valuable."
    )
    set_run_font(ab)

    add_heading(doc, "Background and open materials", 2)
    refs = [
        ("Action for M.E. PRIME research involvement programme", "Open programme page", "https://www.actionforme.org.uk/research-campaigns/our-research-work/prime/"),
        ("Juno Open Health Tools - permanent Zenodo archive", "Open archive", "https://doi.org/10.5281/zenodo.21730086"),
        ("Juno Open Health Tools - source repository", "Open repository", "https://github.com/MarshallBear1/juno-open-health-tools"),
        ("Plain-language guide to the appointment-preparation method", "Read the guide", "https://medium.com/@Marsh30/a-better-way-to-prepare-for-a-health-appointment-when-your-symptoms-are-messy-df9e62d83f5a"),
    ]
    for label, link_text, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=9.5, color=INK, bold=True)
        add_hyperlink(p, link_text, url)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.1
    nr = note.add_run(
        "This concept note is for discussion only. It is not a research protocol, recruitment notice, medical advice or claim of clinical effectiveness. No participant data has been collected."
    )
    set_run_font(nr, size=9, color=MUTED, italic=True)

    props = doc.core_properties
    props.title = "Juno x Action for M.E. - Co-authored data brief concept note"
    props.subject = "Appointment communication and low-burden preparation for people with ME/CFS"
    props.author = "Juno"
    props.keywords = "ME/CFS, patient involvement, appointment preparation, health communication"
    props.comments = "Prepared for partnership discussion. No participant data included."

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
